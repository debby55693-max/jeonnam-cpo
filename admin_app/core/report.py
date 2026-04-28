"""
소상공인 안전물품 지원 관리시스템 - A4 1페이지 리포트

위원회 검토용 점포별 1장 보고서를 Word(.docx) 형태로 생성합니다.
모든 내용이 A4 용지 1페이지 이내에 출력되도록 설계되어 있습니다.
"""

from __future__ import annotations
from io import BytesIO
from typing import Any, Dict, List, Optional
import zipfile
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.scoring import score_breakdown, FIELD_SCORING


# ─────────────────────────────────────────────────
# 공통 헬퍼
# ─────────────────────────────────────────────────

def _s(v: Any) -> str:
    return str(v).strip() if v is not None else ""


def _or(v: Any, fallback: str = "-") -> str:
    s = _s(v)
    return s if s and s not in ("None", "nan") else fallback


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, sides=("top", "bottom", "left", "right"), color="CCCCCC", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), sz)
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _set_no_spacing(para):
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)


def _run(para, text: str, size_pt: float = 8.5, bold: bool = False,
         color_hex: Optional[str] = None, italic: bool = False):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size_pt)
    if color_hex:
        r.font.color.rgb = RGBColor.from_string(color_hex)
    return r


def _set_col_width(table, col_idx: int, width_cm: float):
    """테이블 특정 열 너비 설정"""
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def _mini_heading(cell_or_doc, text: str, bg: str = "1E3A5F"):
    """셀 내부 또는 문서에 미니 섹션 헤딩"""
    if hasattr(cell_or_doc, 'paragraphs'):
        # cell
        p = cell_or_doc.paragraphs[0]
    else:
        p = cell_or_doc.add_paragraph()
    _set_cell_bg(cell_or_doc if hasattr(cell_or_doc, '_tc') else p, bg) if hasattr(cell_or_doc, '_tc') else None
    _set_no_spacing(p)
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    return p


def _compact_table(doc, rows_data: List[tuple], col_widths: List[float],
                   key_bg="EEF2FA", val_bg="FFFFFF", font_size=8.0):
    """
    rows_data: [(key1, val1), ...] 형식 OR
               [(key1, val1, key2, val2), ...] 4열 형식
    """
    ncols = len(col_widths)
    table = doc.add_table(rows=len(rows_data), cols=ncols)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, (cell_text, width) in enumerate(zip(row_data, col_widths)):
            c = row.cells[j]
            c.width = Cm(width)
            p = c.paragraphs[0]
            _set_no_spacing(p)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            is_key = (j % 2 == 0)  # 짝수 인덱스 = 키 셀
            if is_key:
                _set_cell_bg(c, key_bg)
                _run(p, _s(cell_text), size_pt=font_size, bold=True, color_hex="2D3748")
            else:
                _run(p, _or(cell_text), size_pt=font_size, color_hex="111827")
    return table


def _score_bar_text(score: float, max_score: float, width: int = 8) -> str:
    if max_score <= 0:
        return ""
    ratio = min(1.0, score / max_score)
    filled = int(ratio * width)
    return "█" * filled + "░" * (width - filled)


# ─────────────────────────────────────────────────
# 리포트 생성 메인
# ─────────────────────────────────────────────────

def generate_report(
    row: Dict,
    precas_scores: Dict,
    rank: Optional[int] = None,
    total_count: Optional[int] = None,
) -> bytes:
    """
    점포 1건의 A4 1페이지 보고서를 Word bytes로 반환.
    """
    bd = score_breakdown(row, precas_scores)
    doc = Document()

    # ── A4 1페이지 페이지 설정 (좁은 여백) ─────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin   = Cm(1.2)
    section.right_margin  = Cm(1.2)

    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(8.5)
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after  = Pt(0)

    today = datetime.now().strftime("%Y. %m. %d")
    total = bd["total"]

    # ── ① 헤더 배너 ─────────────────────────────────────────
    hdr = doc.add_table(rows=1, cols=2)
    hdr.style = "Table Grid"
    # 왼쪽: 문서 제목
    lc = hdr.rows[0].cells[0]
    lc.width = Cm(11.8)
    _set_cell_bg(lc, "1E3A5F")
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after = Pt(4)
    lr = lp.add_run("소상공인 방범물품 지원 검토 보고서")
    lr.bold = True; lr.font.size = Pt(13)
    lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # 오른쪽: 기관 + 날짜
    rc = hdr.rows[0].cells[1]
    _set_cell_bg(rc, "2B4F8A")
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(4)
    rp.paragraph_format.space_after = Pt(4)
    rr = rp.add_run(f"전라남도경찰청 CPO\n생성: {today}")
    rr.font.size = Pt(7.5); rr.font.color.rgb = RGBColor(0xBF, 0xDB, 0xFE)

    # ── ② 총점 요약 바 ──────────────────────────────────────
    rank_text = (
        f"{rank}위 / {total_count}건  (상위 {round((1 - rank/total_count)*100)}%)"
        if rank and total_count else "-"
    )
    score_table = doc.add_table(rows=1, cols=6)
    score_table.style = "Table Grid"
    sc = score_table.rows[0].cells

    score_data = [
        ("총  점",     f"{total}점",                    "0F2557", True),
        ("순  위",     rank_text,                       "1B3A7A", False),
        ("프리카스",   f"{bd['precas_total']} / 40점",  "1E3A8A", False),
        ("환경조사",   f"{bd['field_total']} / 20점",   "1D4ED8", False),
        ("설  문",     f"{bd['survey_total']} / 30점",  "2563EB", False),
        ("재  량",     f"{bd['discretionary']} / 10점", "3B82F6", False),
    ]
    for i, (label, val, bg, big) in enumerate(score_data):
        c = sc[i]
        _set_cell_bg(c, bg)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(label + "\n")
        r1.font.size = Pt(7)
        r1.font.color.rgb = RGBColor(0xBF, 0xDB, 0xFE)
        r2 = p.add_run(val)
        r2.bold = True
        r2.font.size = Pt(11 if big else 8.5)
        r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 간격
    sp0 = doc.add_paragraph()
    sp0.paragraph_format.space_before = Pt(0)
    sp0.paragraph_format.space_after = Pt(3)

    # ── ③ 2열 메인 레이아웃 (기본정보 | 프리카스+설문) ───────
    main = doc.add_table(rows=1, cols=2)
    main.style = "Table Grid"
    left = main.rows[0].cells[0]
    right = main.rows[0].cells[1]
    left.width = Cm(9.0)
    right.width = Cm(9.6)

    # ── 왼쪽: 기본 정보 ────────────────────────────────
    def _add_to_cell(cell, key: str, val: str, key_bg="EEF2FA"):
        """셀에 단락 하나 추가 (키=회색, 값=흰색)"""
        # 첫 번째 단락 이후는 새 단락 추가
        existing = [p for p in cell.paragraphs if p.text]
        p = cell.add_paragraph() if existing else cell.paragraphs[0]
        _set_no_spacing(p)
        p.paragraph_format.space_before = Pt(0.5)
        p.paragraph_format.space_after = Pt(0.5)
        _run(p, f"{key}: ", size_pt=8.0, bold=True, color_hex="1E3A5F")
        _run(p, _or(val), size_pt=8.0, color_hex="111827")

    def full_addr():
        road = _s(row.get("address_road"))
        detail = _s(row.get("address_detail"))
        full = _s(row.get("full_address"))
        if full:
            return full
        return f"{road} {detail}".strip() or _or(row.get("address_jibun"))

    status_map = {
        "submitted": "접수완료", "under_review": "검토중",
        "docs_requested": "추가서류요청", "reviewed": "검토완료",
        "selection_considered": "선정고려", "selection_consideration": "선정고려",
        "selected": "선정", "excluded": "제외",
    }
    status_label = status_map.get(_s(row.get("current_status")), _s(row.get("current_status")))

    # 왼쪽 셀 헤딩
    lp0 = left.paragraphs[0]
    _set_no_spacing(lp0)
    lp0.paragraph_format.space_before = Pt(2)
    lp0.paragraph_format.space_after = Pt(2)
    _set_cell_bg(left, "FFFFFF")
    lr0 = lp0.add_run("■ 점포 기본 정보")
    lr0.bold = True; lr0.font.size = Pt(8.5)
    lr0.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    _add_to_cell(left, "점포명",   _or(row.get("business_name")))
    _add_to_cell(left, "대표자",   _or(row.get("applicant_name")))
    _add_to_cell(left, "연락처",   _or(row.get("phone")))

    bt = _s(row.get("business_type"))
    bt_other = _s(row.get("business_type_other"))
    biz_display = f"{bt}/{bt_other}" if bt == "기타" and bt_other else bt or bt_other or "-"
    _add_to_cell(left, "업  종",   biz_display)
    _add_to_cell(left, "연매출",   _or(row.get("sales_band")))
    _add_to_cell(left, "주  소",   full_addr())
    _add_to_cell(left, "관할경찰서", _or(row.get("station_label")))
    _add_to_cell(left, "신청일",   _s(row.get("submitted_at"))[:10])
    _add_to_cell(left, "현재상태", status_label)
    _add_to_cell(left, "희망물품", _or(row.get("requested_item")))

    # ── 오른쪽: 프리카스 + 설문 ────────────────────────
    rp0 = right.paragraphs[0]
    _set_no_spacing(rp0)
    rp0.paragraph_format.space_before = Pt(2)
    rp0.paragraph_format.space_after = Pt(2)
    _set_cell_bg(right, "FFFFFF")
    rr0 = rp0.add_run("■ 프리카스 격자 데이터")
    rr0.bold = True; rr0.font.size = Pt(8.5)
    rr0.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    def _add_score_line(cell, label, raw, score, max_score):
        p = cell.add_paragraph()
        _set_no_spacing(p)
        p.paragraph_format.space_before = Pt(0.5)
        p.paragraph_format.space_after = Pt(0.5)
        bar = _score_bar_text(score, max_score, width=7)
        _run(p, f"  {label}: ", size_pt=7.5, bold=True, color_hex="1E3A5F")
        _run(p, f"{_or(raw)}  {bar}  {score}/{max_score}점", size_pt=7.5, color_hex="111827")

    if bd["has_precas"]:
        _add_score_line(right, "112신고", f"{_or(row.get('precas_112_count'))}건",
                        bd['precas_112'], 14)
        _add_score_line(right, "위험도등급", f"{_or(row.get('precas_risk_grade'))}등급",
                        bd['precas_risk'], 12)
        _add_score_line(right, "탄력순찰", f"{_or(row.get('precas_patrol_count'))}회",
                        bd['precas_patrol'], 8)
        _add_score_line(right, "CCTV 대수", f"{_or(row.get('public_cctv_count'))}대",
                        bd.get('precas_cctv', 0), 6)
    else:
        p_np = right.add_paragraph()
        _set_no_spacing(p_np)
        p_np.paragraph_format.space_before = Pt(1)
        _run(p_np, "  프리카스 데이터 미입력 (0점)", size_pt=7.5, color_hex="999999")

    # 구분선
    p_div = right.add_paragraph()
    _set_no_spacing(p_div)
    p_div.paragraph_format.space_before = Pt(2)
    p_div.paragraph_format.space_after = Pt(2)
    r_div = p_div.add_run("─" * 35)
    r_div.font.size = Pt(7); r_div.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)

    # 설문 항목 추가
    p_survtitle = right.add_paragraph()
    _set_no_spacing(p_survtitle)
    p_survtitle.paragraph_format.space_before = Pt(1)
    p_survtitle.paragraph_format.space_after = Pt(1)
    rt = p_survtitle.add_run("■ 신청인 설문")
    rt.bold = True; rt.font.size = Pt(8.5)
    rt.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    survey_items = [
        ("범죄불안경험",  row.get("survey_crime_anxiety")),
        ("야간영업",      row.get("survey_late_night")),
        ("주변환경",      row.get("survey_dark_area")),
        ("단독근무",      row.get("survey_single_worker")),
    ]
    # 2열로 압축 (2개씩 한 줄)
    for i in range(0, len(survey_items), 2):
        p = right.add_paragraph()
        _set_no_spacing(p)
        p.paragraph_format.space_before = Pt(0.5)
        p.paragraph_format.space_after = Pt(0.5)
        k1, v1 = survey_items[i]
        _run(p, f"  {k1}: ", size_pt=7.5, bold=True, color_hex="1E3A5F")
        _run(p, f"{_or(v1)}", size_pt=7.5, color_hex="111827")
        if i + 1 < len(survey_items):
            k2, v2 = survey_items[i + 1]
            _run(p, f"  |  {k2}: ", size_pt=7.5, bold=True, color_hex="1E3A5F")
            _run(p, f"{_or(v2)}", size_pt=7.5, color_hex="111827")

    # CCTV/경비/비상벨
    p_sec = right.add_paragraph()
    _set_no_spacing(p_sec)
    p_sec.paragraph_format.space_before = Pt(0.5)
    p_sec.paragraph_format.space_after = Pt(0.5)
    cctv = "있음" if bool(row.get("has_cctv")) else "없음"
    bell = "있음" if bool(row.get("has_emergency_bell")) else "없음"
    sec_co = "이용" if bool(row.get("uses_security_company")) else "미이용"
    _run(p_sec, "  CCTV: ", size_pt=7.5, bold=True, color_hex="1E3A5F")
    _run(p_sec, f"{cctv}  |  ", size_pt=7.5, color_hex="111827")
    _run(p_sec, "비상벨: ", size_pt=7.5, bold=True, color_hex="1E3A5F")
    _run(p_sec, f"{bell}  |  ", size_pt=7.5, color_hex="111827")
    _run(p_sec, "사설경비: ", size_pt=7.5, bold=True, color_hex="1E3A5F")
    _run(p_sec, sec_co, size_pt=7.5, color_hex="111827")

    # 체감안전도
    p_fs = right.add_paragraph()
    _set_no_spacing(p_fs)
    p_fs.paragraph_format.space_before = Pt(0.5)
    p_fs.paragraph_format.space_after = Pt(0.5)
    _run(p_fs, "  체감안전도: ", size_pt=7.5, bold=True, color_hex="1E3A5F")
    _run(p_fs, f"{bd['felt_safety']}점/5점  |  ", size_pt=7.5, color_hex="111827")
    _run(p_fs, "점포환경설문: ", size_pt=7.5, bold=True, color_hex="1E3A5F")
    _run(p_fs, f"{bd['survey_env']}점/25점", size_pt=7.5, color_hex="111827")

    # ── ④ CPO 현장 환경조사 (4항목씩 2행) ──────────────────
    sp1 = doc.add_paragraph()
    sp1.paragraph_format.space_before = Pt(0)
    sp1.paragraph_format.space_after = Pt(3)

    # 헤딩 행
    env_heading_t = doc.add_table(rows=1, cols=1)
    env_heading_t.style = "Table Grid"
    ehc = env_heading_t.rows[0].cells[0]
    _set_cell_bg(ehc, "1E3A5F")
    ehp = ehc.paragraphs[0]
    ehp.paragraph_format.space_before = Pt(2)
    ehp.paragraph_format.space_after = Pt(2)
    ehr = ehp.add_run(
        f"■ CPO 현장 환경조사  (배점 20점 / 획득 {bd['field_total']}점)"
    )
    ehr.bold = True; ehr.font.size = Pt(8.5)
    ehr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 환경조사 데이터를 4열 테이블로
    field_detail = bd.get("field_detail", [])
    # 4개씩 한 행으로 묶기
    ITEMS_PER_ROW = 4
    n_rows = -(-len(field_detail) // ITEMS_PER_ROW)  # ceiling division
    env_table = doc.add_table(rows=n_rows, cols=ITEMS_PER_ROW * 2)  # key+val 교대
    env_table.style = "Table Grid"

    # 열 너비: 전체 18.6cm → 항목 4개 × (2.0 key + 2.65 val) = 18.6
    KEY_W = 2.3
    VAL_W = 2.325
    for ri in range(n_rows):
        for ci in range(ITEMS_PER_ROW):
            item_idx = ri * ITEMS_PER_ROW + ci
            key_col = ci * 2
            val_col = ci * 2 + 1
            kc = env_table.rows[ri].cells[key_col]
            vc = env_table.rows[ri].cells[val_col]
            kc.width = Cm(KEY_W)
            vc.width = Cm(VAL_W)

            if item_idx < len(field_detail):
                item = field_detail[item_idx]
                _set_cell_bg(kc, "EEF2FA")
                kp = kc.paragraphs[0]
                _set_no_spacing(kp)
                kp.paragraph_format.space_before = Pt(1)
                kp.paragraph_format.space_after = Pt(1)
                _run(kp, item["label"], size_pt=7.5, bold=True, color_hex="2D3748")

                vp = vc.paragraphs[0]
                _set_no_spacing(vp)
                vp.paragraph_format.space_before = Pt(1)
                vp.paragraph_format.space_after = Pt(1)
                val_text = _or(item.get("value"))
                score_text = f"  {item['score']}/{item['max']}점"
                _run(vp, val_text, size_pt=7.5, color_hex="111827")
                _run(vp, score_text, size_pt=7, color_hex="2563EB", bold=True)
            else:
                # 빈 셀
                _set_cell_bg(kc, "F8FAFC")
                _set_cell_bg(vc, "F8FAFC")

    # ── ⑤ 위원회 평가 점수 + 최종 검토 의견 ─────────────────
    sp2 = doc.add_paragraph()
    sp2.paragraph_format.space_before = Pt(0)
    sp2.paragraph_format.space_after = Pt(3)

    bottom_table = doc.add_table(rows=1, cols=2)
    bottom_table.style = "Table Grid"
    disc_cell = bottom_table.rows[0].cells[0]
    opinion_cell = bottom_table.rows[0].cells[1]
    disc_cell.width = Cm(4.8)
    opinion_cell.width = Cm(13.8)

    # 위원회 평가 점수
    _set_cell_bg(disc_cell, "EEF2FA")
    dp = disc_cell.paragraphs[0]
    _set_no_spacing(dp)
    dp.paragraph_format.space_before = Pt(2)
    dp.paragraph_format.space_after = Pt(1)
    _run(dp, "위원회 평가 점수\n", size_pt=7.5, bold=True, color_hex="1E3A5F")
    _run(dp, f"{bd['discretionary']}점 / 10점", size_pt=9, bold=True, color_hex="1E3A5F")

    disc_reason = _or(bd.get("discretionary_reason"))
    if disc_reason != "-":
        dp2 = disc_cell.add_paragraph()
        _set_no_spacing(dp2)
        dp2.paragraph_format.space_before = Pt(2)
        _run(dp2, f"사유: {disc_reason}", size_pt=7, color_hex="475569")

    # 최종 의견
    op = opinion_cell.paragraphs[0]
    _set_no_spacing(op)
    op.paragraph_format.space_before = Pt(2)
    op.paragraph_format.space_after = Pt(1)
    _run(op, "최종 검토 의견  ", size_pt=7.5, bold=True, color_hex="1E3A5F")
    comment = _or(row.get("review_comment"))
    _run(op, comment, size_pt=8, color_hex="111827")

    # 검토 메타
    reviewed_at = _s(row.get("reviewed_at") or "")[:10] or "-"
    op2 = opinion_cell.add_paragraph()
    _set_no_spacing(op2)
    op2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    op2.paragraph_format.space_before = Pt(4)
    _run(
        op2,
        f"검토상태: {status_label}  |  관할: {_or(row.get('station_label'))}  |  검토일: {reviewed_at}",
        size_pt=7.5, color_hex="64748B",
    )

    # ── 저장 ─────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_report_zip(rows: List[Dict], precas_scores: Dict) -> bytes:
    """
    선택된 점포 전체의 보고서를 ZIP으로 묶어 bytes 반환.
    점수 기준 내림차순 정렬 후 순위 부여.
    """
    from core.scoring import compute_total_score

    sorted_rows = sorted(rows, key=lambda r: compute_total_score(r, precas_scores), reverse=True)
    total_count = len(sorted_rows)

    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for rank, row in enumerate(sorted_rows, start=1):
            docx_bytes = generate_report(row, precas_scores, rank=rank, total_count=total_count)
            biz_name = _s(row.get("business_name")) or f"store_{rank}"
            safe_name = "".join(c for c in biz_name if c.isalnum() or c in "가-힣ㄱ-ㅎㅏ-ㅣ _-")
            filename = f"{rank:03d}_{safe_name}.docx"
            zf.writestr(filename, docx_bytes)
    return buf.getvalue()